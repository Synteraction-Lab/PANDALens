import datetime
import json
import os

from src.Utilities.file import is_file_exists, ITEM, DETAILS, TIME, remove_file


def append_json_data(file_name, new_data):
    # Create the file and directories if they don't exist
    if not os.path.exists(file_name):
        os.makedirs(os.path.dirname(file_name), exist_ok=True)
        with open(file_name, "w") as f:
            json.dump({"recordings": []}, f)

    # Load the existing data from the JSON file
    with open(file_name, "r") as f:
        data = json.load(f)

    # Append the new data to the existing data
    data["recordings"].append(new_data)

    # Write the updated data to the JSON file
    with open(file_name, "w") as f:
        json.dump(data, f, indent=4)


def append_csv_data(file_name, data, write_type="a"):
    os.makedirs(os.path.dirname(file_name), exist_ok=True)

    try:
        file = open(file_name, write_type)
        file.write(data)
        file.close()
    except Exception as e:
        print("Failed to write: ", e.__class__)


def record_device_config(file_name, item, details):
    if not is_file_exists(file_name):
        append_csv_data(file_name, f'{ITEM},{DETAILS}\n')

    append_csv_data(file_name, f'{item},{details}\n')


def log_manipulation(file_name, item):
    if not is_file_exists(file_name):
        append_csv_data(file_name, f'{ITEM},{TIME}\n')

    append_csv_data(file_name, f'{item},{datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S")}\n')
    print(f'Logged: {item}')


def generate_output_file(chat_history_path, image_path, title='UbiWriter Output'):
    full_writing = retrieve_full_writing_from_chat_history(chat_history_path)
    moment_summary = retrieve_moment_summary_from_chat_history(chat_history_path)
    images = retrieve_all_images(image_path)
    # get current path's parent directory
    data_path = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    data_path = os.path.join(data_path, "data")
    output_filename = title + ".docx"
    output_path = os.path.join(data_path, output_filename)
    pack_doc(full_writing, moment_summary, images, output_path, title)


def retrieve_full_writing_from_chat_history(chat_history_path):
    import json
    with open(chat_history_path, "r") as f:
        data = json.load(f, strict=False)

    # Get the last json item with "mode" == "full"
    for item in reversed(data["recordings"]):
        try:
            record = json.loads(item["content"], strict=False)
            if record["mode"] == "full":
                return record["response"]["full writing"]
        except:
            pass


def retrieve_moment_summary_from_chat_history(chat_history_path):
    import json
    with open(chat_history_path, "r") as f:
        data = json.load(f, strict=False)

    summary = ""

    # Get the last json item with "mode" == "full"
    for item in reversed(data["recordings"]):
        try:
            record = json.loads(item["content"], strict=False)
            if record["mode"] == "authoring":
                if "summary of new content" in record["response"]:
                    summary += record["response"]["summary of new content"] + "\n\n"
        except:
            pass

    return summary + "--------------------------\n\n"


def retrieve_all_images(image_path):
    # return all images in the folder
    import os
    # return all the png and jpg files in the folder
    return [os.path.join(image_path, f) for f in os.listdir(image_path) if
            os.path.isfile(os.path.join(image_path, f)) and
            (f.endswith(".png") or f.endswith(".jpg"))]


def pack_doc(full_writing, moment_summary, images, output_path, title):
    from docx import Document
    from docx.shared import Inches
    document = Document()
    document.add_heading(title, 0)

    paragraph = document.add_paragraph()
    for image in images:
        run = paragraph.add_run()
        run.add_picture(image, width=Inches(1.25))

    document.add_paragraph(f'{moment_summary}\n')
    document.add_paragraph(f'{full_writing}\n')

    document.save(output_path)
    print("Doc is saved to:", output_path)


def remove_file_async(file_name):
    import threading
    threading.Thread(target=remove_file, args=(file_name,)).start()

